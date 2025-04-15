import json
import csv
import os
from datetime import datetime
from docx import Document
from openpyxl import Workbook
from rich.console import Console

# Variáveis globais
arq_json = 'alunos.json'
arq_csv = 'alunos.csv'

# Classe Aluno para armazenar os dados
class Aluno:
    def __init__(self, nome, matricula, data_nascimento, notas, faltas):
        self.nome = nome
        self.matricula = matricula
        self.data_nascimento = data_nascimento
        self.notas = notas
        self.faltas = faltas
    
    @property
    def media(self):
        return sum(self.notas) / len(self.notas)
    
    @property
    def status(self):
        if self.media >= 7 and self.faltas <= 25:
            return 'Aprovado'
        else:
            return 'Reprovado'

# Função para carregar os alunos
def carregar_alunos():
    alunos = []
    if os.path.exists(arq_json):
        with open(arq_json, 'r') as f:
            alunos_data = json.load(f)
            for aluno_data in alunos_data:
                aluno = Aluno(
                    aluno_data['nome'],
                    aluno_data['matricula'],
                    aluno_data['data_nascimento'],
                    aluno_data['notas'],
                    aluno_data['faltas']
                )
                alunos.append(aluno)
    return alunos

# Função para salvar os alunos
def salvar_alunos(alunos):
    alunos_data = [{
        'nome': aluno.nome,
        'matricula': aluno.matricula,
        'data_nascimento': aluno.data_nascimento,
        'notas': aluno.notas,
        'faltas': aluno.faltas
    } for aluno in alunos]
    
    with open(arq_json, 'w') as f:
        json.dump(alunos_data, f, indent=4)

    # Para salvar como CSV
    with open(arq_csv, 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['nome', 'matricula', 'data_nascimento', 'notas', 'faltas'])
        for aluno in alunos:
            writer.writerow([aluno.nome, aluno.matricula, aluno.data_nascimento, ','.join(map(str, aluno.notas)), aluno.faltas])

# Função para adicionar aluno
def adicionar_aluno(alunos):
    nome = input("Nome completo: ")
    matricula = input("Matrícula: ")
    data_nascimento = input("Data de nascimento (dd/mm/yyyy): ")
    notas = list(map(float, input("Notas (separadas por espaço): ").split()))
    faltas = int(input("Total de faltas: "))
    
    aluno = Aluno(nome, matricula, data_nascimento, notas, faltas)
    alunos.append(aluno)
    salvar_alunos(alunos)
    print("Aluno adicionado com sucesso!")

# Função para remover aluno
def remover_aluno(alunos):
    matricula = input("Digite a matrícula do aluno a ser removido: ")
    aluno_a_remover = None
    for aluno in alunos:
        if aluno.matricula == matricula:
            aluno_a_remover = aluno
            break
    
    if aluno_a_remover:
        alunos.remove(aluno_a_remover)
        salvar_alunos(alunos)
        print(f"Aluno {aluno_a_remover.nome} removido com sucesso!")
    else:
        print("Aluno não encontrado!")

# Função para listar alunos
def listar_alunos(alunos):
    print(f"{'Nome':<30} {'Matrícula':<15} {'Notas':<25} {'Faltas':<10}")
    print("-" * 80)
    for aluno in alunos:
        print(f"{aluno.nome:<30} {aluno.matricula:<15} {', '.join(map(str, aluno.notas)):<25} {aluno.faltas:<10}")

# Função para gerar relatório individual em Word
def gerar_relatorio_individual(aluno):
    doc = Document()
    doc.add_heading(f"Relatório de {aluno.nome}", 0)
    
    doc.add_paragraph(f"Matrícula: {aluno.matricula}")
    doc.add_paragraph(f"Data de Nascimento: {aluno.data_nascimento}")
    doc.add_paragraph(f"Notas: {', '.join(map(str, aluno.notas))}")
    doc.add_paragraph(f"Média: {aluno.media:.2f}")
    doc.add_paragraph(f"Total de faltas: {aluno.faltas}")
    doc.add_paragraph(f"Status: {aluno.status}")
    
    nome_arquivo = f"relatorio_{aluno.matricula}.docx"
    doc.save(nome_arquivo)
    print(f"Relatório gerado: {nome_arquivo}")

# Função para gerar relatório em grupo em Word
def gerar_relatorio_grupo(alunos):
    doc = Document()
    doc.add_heading("Relatório de Alunos", 0)
    
    for aluno in alunos:
        doc.add_paragraph(f"\nAluno: {aluno.nome}")
        doc.add_paragraph(f"Matrícula: {aluno.matricula}")
        doc.add_paragraph(f"Notas: {', '.join(map(str, aluno.notas))}")
        doc.add_paragraph(f"Média: {aluno.media:.2f}")
        doc.add_paragraph(f"Total de faltas: {aluno.faltas}")
        doc.add_paragraph(f"Status: {aluno.status}")
    
    doc.save("relatorio_grupo.docx")
    print("Relatório de grupo gerado: relatorio_grupo.docx")

# Função para exportar dados para Excel
def exportar_para_excel(alunos):
    wb = Workbook()
    ws = wb.active
    ws.append(['Nome', 'Matrícula', 'Data de Nascimento', 'Notas', 'Média', 'Faltas', 'Status'])
    
    for aluno in alunos:
        ws.append([aluno.nome, aluno.matricula, aluno.data_nascimento, ', '.join(map(str, aluno.notas)), aluno.media, aluno.faltas, aluno.status])
    
    wb.save("relatorio_alunos.xlsx")
    print("Relatório exportado para Excel: relatorio_alunos.xlsx")

# Função principal para rodar o menu
def menu():
    alunos = carregar_alunos()
    console = Console()

    while True:
        console.print("\n[bold cyan]Sistema de Gestão Escolar[/bold cyan]")
        print("1. Adicionar Aluno")
        print("2. Remover Aluno")
        print("3. Listar Alunos")
        print("4. Gerar Relatório Individual")
        print("5. Gerar Relatório em Grupo")
        print("6. Exportar para Excel")
        print("7. Sair")
        
        opcao = input("Escolha uma opção: ")
        
        if opcao == '1':
            adicionar_aluno(alunos)
        elif opcao == '2':
            remover_aluno(alunos)
        elif opcao == '3':
            listar_alunos(alunos)
        elif opcao == '4':
            matricula = input("Digite a matrícula do aluno: ")
            aluno = next((a for a in alunos if a.matricula == matricula), None)
            if aluno:
                gerar_relatorio_individual(aluno)
            else:
                print("Aluno não encontrado!")
        elif opcao == '5':
            gerar_relatorio_grupo(alunos)
        elif opcao == '6':
            exportar_para_excel(alunos)
        elif opcao == '7':
            print("Saindo...")
            break
        else:
            print("Opção inválida!")

if __name__ == '__main__':
    menu()